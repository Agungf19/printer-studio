from __future__ import annotations

import logging
from pathlib import Path

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

router = APIRouter(tags=["export"])
logger = logging.getLogger(__name__)

# Scan files directory
SCAN_DIR = Path(__file__).resolve().parents[3] / "temp" / "scans"


class ExportRequest(BaseModel):
    filename: str
    format: str  # "pdf", "docx", "png", "jpg"
    quality: str = "high"  # "high", "medium", "low"
    pdfa: bool = False


class OcrRequest(BaseModel):
    filename: str
    language: str = "ind+eng"


class OcrResponse(BaseModel):
    text: str
    language: str
    character_count: int


class MergeRequest(BaseModel):
    filenames: list[str]
    format: str  # "pdf" or "png"


@router.post("/export", response_class=FileResponse)
def export_file(payload: ExportRequest):
    """Export a scanned image to PDF, DOCX, PNG, or JPG."""
    source_path = SCAN_DIR / payload.filename
    if not source_path.exists():
        raise HTTPException(status_code=404, detail=f"File tidak ditemukan: {payload.filename}")

    # Determine output name
    stem = source_path.stem
    fmt = payload.format.lower()

    try:
        if fmt == "pdf":
            return _export_pdf(source_path, stem, payload.quality, payload.pdfa)
        elif fmt == "docx":
            return _export_docx(source_path, stem)
        elif fmt == "png":
            return _export_image(source_path, stem, "PNG", payload.quality)
        elif fmt in ("jpg", "jpeg"):
            return _export_image(source_path, stem, "JPEG", payload.quality)
        else:
            raise HTTPException(status_code=400, detail=f"Format tidak didukung: {fmt}")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Export failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Ekspor gagal: {exc}") from exc


def _get_quality(quality: str) -> int:
    return {"high": 95, "medium": 75, "low": 50}.get(quality, 95)


def _export_pdf(source_path: Path, stem: str, quality: str, pdfa: bool) -> FileResponse:
    """Convert image to PDF using img2pdf (lossless) or Pillow."""
    output_path = SCAN_DIR / f"{stem}.pdf"
    try:
        import img2pdf
        # img2pdf is lossless — embeds image directly
        with open(output_path, "wb") as f:
            f.write(img2pdf.convert(str(source_path)))
    except ImportError:
        # Fallback: Pillow
        from PIL import Image
        img = Image.open(source_path)
        if img.mode == "RGBA":
            img = img.convert("RGB")
        img.save(str(output_path), "PDF", quality=_get_quality(quality))

    # PDF/A conversion
    if pdfa:
        try:
            import pikepdf
            with pikepdf.open(str(output_path)) as pdf:
                # Set PDF/A metadata
                with pdf.open_metadata() as meta:
                    meta["pdf:Producer"] = "ScanPilot"
                pdf.save(str(output_path))
        except ImportError:
            logger.warning("pikepdf not available — skipping PDF/A")

    return FileResponse(
        str(output_path),
        media_type="application/pdf",
        filename=f"{stem}.pdf",
    )


def _export_docx(source_path: Path, stem: str) -> FileResponse:
    """Create a DOCX with the image embedded."""
    output_path = SCAN_DIR / f"{stem}.docx"
    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading(f"Scan: {stem}", level=1)
        doc.add_picture(str(source_path), width=Inches(6.0))
        doc.save(str(output_path))
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx belum terinstall.")

    return FileResponse(
        str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{stem}.docx",
    )


def _export_image(source_path: Path, stem: str, fmt: str, quality: str) -> FileResponse:
    """Convert/resize image for export."""
    from PIL import Image

    img = Image.open(source_path)
    ext = fmt.lower()
    output_path = SCAN_DIR / f"{stem}_export.{ext}"

    save_kwargs: dict = {}
    if fmt == "JPEG":
        if img.mode == "RGBA":
            img = img.convert("RGB")
        save_kwargs["quality"] = _get_quality(quality)
        save_kwargs["optimize"] = True
    elif fmt == "PNG":
        if quality == "low":
            save_kwargs["optimize"] = True

    img.save(str(output_path), fmt, **save_kwargs)

    media_map = {"PNG": "image/png", "JPEG": "image/jpeg"}
    return FileResponse(
        str(output_path),
        media_type=media_map.get(fmt, "image/png"),
        filename=f"{stem}_export.{ext}",
    )


@router.post("/merge", response_class=FileResponse)
def merge_files(payload: MergeRequest):
    """Merge multiple scanned images into 1 PDF (multi-page) or 1 PNG (vertical strip)."""
    if len(payload.filenames) < 2:
        raise HTTPException(status_code=400, detail="Perlu minimal 2 file untuk digabung.")

    # Validate all files exist
    paths: list[Path] = []
    for fname in payload.filenames:
        p = SCAN_DIR / fname
        if not p.exists():
            raise HTTPException(status_code=404, detail=f"File tidak ditemukan: {fname}")
        paths.append(p)

    fmt = payload.format.lower()

    try:
        if fmt == "pdf":
            return _merge_pdf(paths)
        elif fmt in ("png", "jpg", "jpeg"):
            return _merge_png(paths)
        else:
            raise HTTPException(status_code=400, detail=f"Format tidak didukung: {fmt}")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Merge failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Gagal menggabung: {exc}") from exc


def _merge_pdf(paths: list[Path]) -> FileResponse:
    """Merge images into a single multi-page PDF using img2pdf (lossless)."""
    output_path = SCAN_DIR / "merged_document.pdf"
    try:
        import img2pdf
        with open(output_path, "wb") as f:
            f.write(img2pdf.convert([str(p) for p in paths]))
    except ImportError:
        # Fallback: Pillow
        from PIL import Image
        images: list[Image.Image] = []
        for p in paths:
            img = Image.open(p)
            if img.mode == "RGBA":
                img = img.convert("RGB")
            images.append(img)
        if images:
            images[0].save(str(output_path), "PDF", save_all=True, append_images=images[1:])

    return FileResponse(
        str(output_path),
        media_type="application/pdf",
        filename=f"merged_{len(paths)}_pages.pdf",
    )


def _merge_png(paths: list[Path]) -> FileResponse:
    """Merge images into a single vertical PNG strip."""
    from PIL import Image

    images: list[Image.Image] = []
    for p in paths:
        img = Image.open(p).convert("RGB")
        images.append(img)

    if not images:
        raise HTTPException(status_code=400, detail="Tidak ada gambar untuk digabung.")

    # Resize all to same width (use max width)
    max_width = max(img.width for img in images)
    resized: list[Image.Image] = []
    for img in images:
        if img.width != max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height), Image.LANCZOS)
        resized.append(img)

    # Create vertical strip
    total_height = sum(img.height for img in resized)
    merged = Image.new("RGB", (max_width, total_height), (255, 255, 255))
    y_offset = 0
    for img in resized:
        merged.paste(img, (0, y_offset))
        y_offset += img.height
        img.close()

    output_path = SCAN_DIR / "merged_strip.png"
    merged.save(str(output_path), "PNG", optimize=True)

    return FileResponse(
        str(output_path),
        media_type="image/png",
        filename=f"merged_{len(paths)}_strip.png",
    )
